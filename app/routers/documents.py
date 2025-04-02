from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status
from sqlalchemy.orm import Session
from typing import List
import secrets
import os  # Import os for file path handling
from fastapi.responses import FileResponse  # Import FileResponse
from app.database import get_db
from app.models.document import Document
from app.schemas.document import DocumentResponse
from app.services.auth import get_current_user
from app.models.user import User
from app.config import settings  # Import settings to access environment variables
from PIL import Image, ImageDraw, ImageFont
import PyPDF2
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from datetime import datetime  # Import datetime
import io  # Import io for BytesIO
import docx  # Make sure to install python-docx

router = APIRouter()

@router.post("/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # Generate unique access token
    access_token = secrets.token_urlsafe(32)
    
    # Define the file path
    file_path = f"documents/{access_token}_{file.filename}"
    signed_file_path = f"documents/signed_{access_token}_{file.filename}"

    # Ensure the documents directory exists
    os.makedirs(os.path.dirname(file_path), exist_ok=True)

    # Save the uploaded file
    with open(file_path, "wb") as f:
        f.write(await file.read())

    # Check if the uploaded file is a PDF
    if not file.filename.lower().endswith('.pdf'):
        # Convert to PDF using reportlab
        pdf_file_path = f"documents/{access_token}_{file.filename.rsplit('.', 1)[0]}.pdf"
        c = canvas.Canvas(pdf_file_path, pagesize=letter)

        # Read the content of the uploaded file based on its type
        if file.filename.lower().endswith('.txt'):
            content = await file.read()  # Read the content of the uploaded file
            text_content = content.decode('utf-8')  # Assuming the file is a text file
            c.drawString(100, 750, f"Uploaded Document: {file.filename}")
            c.drawString(100, 730, text_content)  # Write the actual content
        elif file.filename.lower().endswith('.docx'):
            doc = docx.Document(io.BytesIO(await file.read()))
            text_content = "\n".join([para.text for para in doc.paragraphs])
            c.drawString(100, 750, f"Uploaded Document: {file.filename}")
            c.drawString(100, 730, text_content)  # Write the actual content
        else:
            c.drawString(100, 750, f"Uploaded Document: {file.filename}")
            c.drawString(100, 730, "Unsupported file type for content extraction.")  # Handle unsupported types

        c.save()

        # Update the file path to the new PDF
        file_path = pdf_file_path

    # Create document record
    db_document = Document(
        filename=file.filename,
        file_path=file_path,
        signed_file_path=signed_file_path,
        access_token=access_token,
        user_id=current_user.id
    )
    db.add(db_document)
    db.commit()
    db.refresh(db_document)

    # Construct file links
    file_link = f"{settings.BASE_URL}/{db_document.file_path}"
    signed_link = f"{settings.BASE_URL}/{db_document.signed_file_path}" if db_document.signed_at else None

    return DocumentResponse(
        id=db_document.id,
        filename=db_document.filename,
        access_token=db_document.access_token,
        created_at=db_document.created_at,
        signed_at=db_document.signed_at,
        file_link=file_link,
        signed_link=signed_link
    )

@router.get("/list", response_model=List[DocumentResponse])
async def list_documents(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    documents = db.query(Document).filter(Document.user_id == current_user.id).all()
    
    # Construct file links for each document
    for document in documents:
        document.file_link = f"{settings.BASE_URL}/{document.file_path}"
        document.signed_link = f"{settings.BASE_URL}/{document.signed_file_path}" if document.signed_at else None

    return documents

@router.get("/metadata/{access_token}", response_model=DocumentResponse)
async def get_document(
    access_token: str,
    db: Session = Depends(get_db)
):
    document = db.query(Document).filter(Document.access_token == access_token).first()
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    # Construct the file link and signed link using BASE_URL from environment variables
    file_link = f"{settings.BASE_URL}/{document.file_path}"  # Use BASE_URL
    signed_link = f"{settings.BASE_URL}/{document.signed_file_path}"  # Use BASE_URL

    return DocumentResponse(
        id=document.id,
        filename=document.filename,
        access_token=document.access_token,
        created_at=document.created_at,
        signed_at=document.signed_at,
        file_link=file_link,
        signed_link=signed_link
    )

@router.get("/view/{access_token}")
async def view_document(
    access_token: str,
    db: Session = Depends(get_db)
):
    document = db.query(Document).filter(Document.access_token == access_token).first()
    
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    # Check if the document is signed
    if document.signed_at:
        # Serve the signed document
        return FileResponse(document.signed_file_path)  # Directly return the signed document
    else:
        # Serve the original document
        return FileResponse(document.file_path)  # Directly return the original document

